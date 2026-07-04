"""
BVRITH Website Scraper — creates a comprehensive, well-structured DOCX knowledge base.
"""
import os, re, time, logging, json
from typing import List, Dict, Tuple, Set
from urllib.parse import urljoin, urlparse
import requests
from bs4 import BeautifulSoup
from docx import Document

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(message)s")
logger = logging.getLogger(__name__)

BASE_URL = "https://bvrithyderabad.edu.in/"
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "data", "bvrit_info.docx")

# Section mapping for URLs
SECTION_MAP = {
    "about": "About", "vision": "About", "mission": "About", "principal": "About",
    "management": "About", "governing": "About",
    "cse": "Departments", "computer": "Departments", "electronics": "Departments",
    "electrical": "Departments", "mechanical": "Departments", "civil": "Departments",
    "humanities": "Departments", "basic-science": "Departments", "physics": "Departments",
    "chemistry": "Departments", "mathematics": "Departments", "english": "Departments",
    "ai-ml": "Departments", "artificial": "Departments", "data-science": "Departments",
    "cyber": "Departments", "information": "Departments", "it": "Departments",
    "ece": "Departments", "eee": "Departments", "mech": "Departments",
    "admission": "Admissions", "fee": "Fee Structure",
    "placement": "Placements", "training": "Placements", "recruiter": "Placements",
    "internship": "Placements", "employability": "Placements",
    "facilit": "Facilities", "library": "Facilities", "lab": "Facilities",
    "sport": "Facilities", "hostel": "Facilities", "transport": "Facilities",
    "cafeteria": "Facilities", "canteen": "Facilities", "gym": "Facilities",
    "medical": "Facilities", "health": "Facilities",
    "faculty": "Faculty", "staff": "Faculty", "fdp": "Faculty",
    "contact": "Contact",
    "research": "Research", "rd": "Research", "patent": "Research",
    "publication": "Research", "ph.d": "Research",
    "student": "Student Life", "club": "Student Life", "cultural": "Student Life",
    "nss": "Student Life", "ncc": "Student Life",
    "alumni": "Alumni",
    "award": "Awards", "achievement": "Awards",
    "news": "News", "event": "News", "synergia": "News",
    "nirf": "NIRF", "ranking": "NIRF",
    "naac": "NAAC", "accreditation": "NAAC", "nba": "NAAC",
    "iic": "IIC", "innovation": "IIC", "startup": "IIC",
    "entrepreneurship": "IIC", "incubation": "IIC",
    "mou": "IIC", "collaboration": "IIC", "tie-up": "IIC",
    "industry": "IIC",
    "committee": "Committees", "grievance": "Committees",
    "anti-ragging": "Committees", "women": "Committees",
    "sexual": "Committees", "internal": "Committees",
    "complaints": "Committees", "sc-st": "Committees",
    "obc": "Committees", "minority": "Committees",
    "equal": "Committees", "redressal": "Committees",
    "discipline": "Committees", "examination": "Committees",
    "ragging": "Committees", "harass": "Committees",
    "cell": "Committees",
    "gallery": "Gallery", "photo": "Gallery",
    "video": "Gallery", "media": "Gallery",
    "seminar": "Events", "workshop": "Events", "conference": "Events",
    "symposium": "Events", "guest": "Events", "lecture": "Events",
    "talk": "Events", "webinar": "Events",
    "synergia": "Events", "fiesta": "Events"
}

def discover_pages(base_url: str) -> List[Tuple[str, str]]:
    """Discover all internal pages from the homepage."""
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    try:
        resp = requests.get(base_url, headers=headers, timeout=30, verify=False)
        resp.raise_for_status()
    except Exception as e:
        logger.error(f"Failed to fetch homepage: {e}")
        return []

    soup = BeautifulSoup(resp.text, "html.parser")
    links: Set[str] = set()
    base_domain = urlparse(base_url).netloc

    for a in soup.find_all("a", href=True):
        href = a["href"]
        full_url = urljoin(base_url, href)
        parsed = urlparse(full_url)
        if parsed.netloc != base_domain and not parsed.netloc.endswith("." + base_domain):
            continue
        if parsed.fragment or parsed.query:
            continue
        if any(ext in parsed.path.lower() for ext in [".pdf", ".doc", ".zip", ".jpg", ".png", ".mp4", ".mp3"]):
            continue
        if parsed.path in ["", "/", "#"]:
            continue
        links.add(full_url)

    # Add important paths
    for path in ["/about/", "/admissions/", "/fee-structure/", "/placements/", "/facilities/", "/contact/", "/faculty/", "/departments/", "/academics/"]:
        links.add(urljoin(base_url, path))

    # Sort and map to sections
    result = []
    for url in sorted(links):
        full_url = url.rstrip("/") + "/" if not url.endswith("/") else url
        url_lower = full_url.lower()
        section = "About"
        for keyword, sec in SECTION_MAP.items():
            if keyword in url_lower:
                section = sec
                break
        result.append((section, full_url))

    logger.info(f"Discovered {len(result)} pages")
    return result

def extract_text(url: str) -> Tuple[str, str]:
    """Fetch a URL and extract clean text."""
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    try:
        resp = requests.get(url, headers=headers, timeout=30, verify=False)
        resp.raise_for_status()
    except Exception as e:
        return ("", "")

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup.find_all(["script", "style", "nav", "footer", "header", "aside", "noscript", "iframe"]):
        tag.decompose()

    title = soup.title.get_text(strip=True) if soup.title else ""

    # Find main content
    text = ""
    for selector in ["main", "article", ".entry-content", ".post-content", ".content",
                      "#content", ".site-content", ".page-content",
                      ".et_pb_section", ".elementor-section", ".fusion-content",
                      ".ast-row", ".ast-container", "#primary", "#main",
                      "div[class*='content']", "div[class*='entry']"]:
        container = soup.select_one(selector)
        if container:
            text = container.get_text(separator="\n", strip=True)
            if len(text) > 200:
                break

    if len(text) < 100 and soup.body:
        text = soup.body.get_text(separator="\n", strip=True)

    # Clean
    lines = []
    for line in text.split("\n"):
        line = line.strip()
        if len(line) < 3: continue
        if re.search(r'copyright|©|all rights reserved', line, re.I): continue
        if re.match(r'^https?://', line): continue
        lines.append(line)

    return title, "\n".join(lines)

def scrape_site() -> Dict[str, List[str]]:
    """Scrape the entire site and organize by section."""
    pages = discover_pages(BASE_URL)
    sections: Dict[str, List[str]] = {}

    # Important pages to scrape first
    important_pages = [
        ("About", BASE_URL),
        ("About", urljoin(BASE_URL, "/about/")),
        ("About", urljoin(BASE_URL, "/about/vision-mission/")),
        ("About", urljoin(BASE_URL, "/about/salient-features/")),
        ("About", urljoin(BASE_URL, "/about/principal-message/")),
        ("About", urljoin(BASE_URL, "/about/governing-body/")),
        ("About", urljoin(BASE_URL, "/about/management/")),
        ("Departments", urljoin(BASE_URL, "/departments/")),
        ("Departments", urljoin(BASE_URL, "/academics/")),
        ("Departments", urljoin(BASE_URL, "/computer-science-and-engineering/")),
        ("Departments", urljoin(BASE_URL, "/electronics-and-communication-engineering/")),
        ("Departments", urljoin(BASE_URL, "/electrical-and-electronics-engineering/")),
        ("Departments", urljoin(BASE_URL, "/mechanical-engineering/")),
        ("Departments", urljoin(BASE_URL, "/civil-engineering/")),
        ("Departments", urljoin(BASE_URL, "/humanities-and-sciences/")),
        ("Departments", urljoin(BASE_URL, "/cse-artificial-intelligence-and-machine-learning/")),
        ("Departments", urljoin(BASE_URL, "/data-science/")),
        ("Departments", urljoin(BASE_URL, "/cyber-security/")),
        ("Departments", urljoin(BASE_URL, "/information-technology/")),
        ("Admissions", urljoin(BASE_URL, "/admissions/")),
        ("Admissions", urljoin(BASE_URL, "/admission/admission-process/")),
        ("Admissions", urljoin(BASE_URL, "/admission/eamcet-ranks/")),
        ("Admissions", urljoin(BASE_URL, "/admission/documents-to-submit/")),
        ("Admissions", urljoin(BASE_URL, "/admission/b-category/")),
        ("Admissions", urljoin(BASE_URL, "/admission/fee-details/")),
        ("Admissions", urljoin(BASE_URL, "/admission/hostel/")),
        ("Admissions", urljoin(BASE_URL, "/admission/transportation/")),
        ("Admissions", urljoin(BASE_URL, "/admission/intake-of-courses/")),
        ("Fee Structure", urljoin(BASE_URL, "/fee-structure/")),
        ("Placements", urljoin(BASE_URL, "/placements/")),
        ("Placements", urljoin(BASE_URL, "/placements/placement-details/")),
        ("Placements", urljoin(BASE_URL, "/placements/training-and-placement-cell/")),
        ("Placements", urljoin(BASE_URL, "/placements/training-placement-process/")),
        ("Placements", urljoin(BASE_URL, "/placements/training-and-placement-team/")),
        ("Placements", urljoin(BASE_URL, "/placements/employability-skills/")),
        ("Placements", urljoin(BASE_URL, "/placements/internships/")),
        ("Placements", urljoin(BASE_URL, "/placements/testimonials/")),
        ("Placements", urljoin(BASE_URL, "/placements/our-recruiters/")),
        ("Facilities", urljoin(BASE_URL, "/facilities/")),
        ("Facilities", urljoin(BASE_URL, "/library/")),
        ("Facilities", urljoin(BASE_URL, "/pcs-facilities/")),
        ("Facilities", urljoin(BASE_URL, "/food-and-cafetaria/")),
        ("Facilities", urljoin(BASE_URL, "/gym/")),
        ("Facilities", urljoin(BASE_URL, "/security/")),
        ("Facilities", urljoin(BASE_URL, "/temple/")),
        ("Facilities", urljoin(BASE_URL, "/yoga/")),
        ("Facilities", urljoin(BASE_URL, "/entry-exit-system/")),
        ("Facilities", urljoin(BASE_URL, "/differentiators/assistive-technology-lab/")),
        ("Facilities", urljoin(BASE_URL, "/differentiators/drone-technology-laboratory/")),
        ("Facilities", urljoin(BASE_URL, "/differentiators/iot-maker-space/")),
        ("Facilities", urljoin(BASE_URL, "/research/research-facility/")),
        ("Faculty", urljoin(BASE_URL, "/faculty/")),
        ("Faculty", urljoin(BASE_URL, "/research/faculty-domain-areas/")),
        ("Faculty", urljoin(BASE_URL, "/research/faculty-publications/")),
        ("Faculty", urljoin(BASE_URL, "/research/faculty-as-ph-d-supervisors/")),
        ("Contact", urljoin(BASE_URL, "/contact/")),
        ("Contact", urljoin(BASE_URL, "/contact-us/")),
    ]

    seen = set()
    for sec, url in important_pages + pages:
        if url in seen: continue
        seen.add(url)

        title, text = extract_text(url)
        if not text or len(text) < 50:
            continue

        if sec not in sections:
            sections[sec] = []
        sections[sec].append(f"--- {sec}: {title} ---\n{text}\n")
        logger.info(f"  [{sec}] {len(text)} chars")

        time.sleep(0.2)

    return sections

def create_docx(sections: Dict[str, List[str]], output_path: str) -> None:
    """Create a well-structured DOCX with clear section headers."""
    doc = Document()
    doc.add_heading("BVRITH College Knowledge Base", 0)
    doc.add_paragraph("Source: https://bvrithyderabad.edu.in/")
    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_page_break()

    # Order sections
    section_order = [
        "About", "Departments", "Admissions", "Fee Structure", "Placements",
        "Facilities", "Faculty", "Contact", "Research", "Student Life",
        "Alumni", "Awards", "NIRF", "NAAC", "IIC", "Committees",
        "News", "Events", "Gallery"
    ]

    for sec_name in section_order:
        if sec_name not in sections:
            continue
        contents = sections[sec_name]
        doc.add_heading(sec_name, level=1)
        seen_lines = set()
        for content in contents:
            for line in content.split("\n"):
                line = line.strip()
                if not line or line in seen_lines: continue
                seen_lines.add(line)
                if line.startswith("---"):
                    doc.add_heading(line.strip("- "), level=2)
                else:
                    doc.add_paragraph(line)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    logger.info(f"DOCX saved: {output_path}")

def main():
    logger.info("=" * 50)
    logger.info("BVRITH Website Scraper")
    logger.info("=" * 50)

    sections = scrape_site()

    total_chars = sum(len(t) for contents in sections.values() for t in contents)
    logger.info(f"\nScraped {len(sections)} sections, {total_chars} total chars")
    for sec, contents in sections.items():
        chars = sum(len(c) for c in contents)
        logger.info(f"  {sec}: {len(contents)} pages, {chars} chars")

    create_docx(sections, OUTPUT_PATH)
    logger.info(f"DONE! Output: {OUTPUT_PATH}")

if __name__ == "__main__":
    import urllib3; urllib3.disable_warnings()
    main()